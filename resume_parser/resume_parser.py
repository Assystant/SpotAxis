from __future__ import absolute_import, print_function
import os
import re
import json
import logging
from docx import Document
from bs4 import BeautifulSoup
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from striprtf.striprtf import rtf_to_text
from xml.dom.minidom import parseString
from unidecode import unidecode
import dicttoxml


logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

SECTION_HEADERS = {
    "education": ["education", "academic", "academics"],
    "skills": ["skills", "technical skills", "technologies", "expertise"],
    "experience": ["experience", "employment", "work", "professional experience"],
    "projects": ["projects", "personal projects"],
    "summary": ["summary", "objective", "profile"],
    "certifications": ["certifications", "licenses"]
}

omit_list = ['High School', 'Senior Secondary School', 'Secondary School', 'Higher Secondary School', 'School']
school_college_keyword_list = [
    'COLLEGE', 'College', 'INSTITUTE', 'Institute', 'SCHOOL', 'School', 'UNIVERSITY', 'University'
]
regex_expression_date = r'(?:19[7-9]\d|2\d{3}) *-? *((?:19[7-9]\d|2\d{3})|Present|current|Current|present|ongoing|Ongoing)?'
regex_expression_school_college = 'COLLEGE|College|INSTITUTE|Institute|SCHOOL|School|UNIVERSITY|University'
re_for_alpha = r'[a-zA-Z\.]'

module_dir = os.path.dirname(__file__)
def load_lines(filename):
    try:
        with open(os.path.join(module_dir, filename), encoding='utf-8') as f:
            return [line.strip() for line in f.readlines()]
    except Exception:
        return []

def load_json(filename):
    try:
        with open(os.path.join(module_dir, filename), encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}

course_corpus_list = load_lines('courses_corpus.txt')
female_list = load_lines('indian_female_list.txt')
male_list = load_lines('indian_male_list.txt')
skill_json_data = load_json('skill_data_json.json')

def normalize_text(text):
    return re.sub(r'\s+', ' ', unidecode(text))

def match_section(line):
    for canonical, aliases in SECTION_HEADERS.items():
        if any(alias.lower() in line.lower() for alias in aliases):
            return canonical
    return None

def extract_links(text):
    return re.findall(r'https?://\S+', text)

def safe_extract(func, fallback):
    try:
        return func()
    except Exception as e:
        logging.error(f"Error in {func.__name__}: {e}")
        return fallback

def get_email(file_content):
    try:
        return re.findall(r'[\w\.-]+@[\w\.-]+', file_content)
    except Exception as e:
        logging.error(f"Email extraction failed: {e}")
        return []

def get_phone_number(file_content):
    try:
        phones = re.findall(r'\+?\d[\d\s\-\(\)]{9,}', file_content)
        return [re.sub(r'[^\d+]', '', phone) for phone in phones if len(re.sub(r'[^\d]', '', phone)) >= 10]
    except Exception as e:
        logging.error(f"Phone number extraction failed: {e}")
        return []

def get_name(file_content):
    try:
        name_user_set = set()
        local_name_block_string = ''
        for line in file_content.split('\n'):
            if any(word in line for word in ['SKILL', 'Skill', 'OBJECTIVE', 'Objective', 'EXPERIENCE', 'Experience']):
                break
            if line.strip():
                local_name_block_string += line.strip() + '\n'
        local_list_name_block = [_f for _f in local_name_block_string.split('\n') if _f]
        for item in local_list_name_block:
            if 'NAME' in item or 'Name' in item:
                name_user_set.add(item.replace('Name', '').replace('NAME', '').strip())
            if bool(re.search(r'\d', item)):
                continue
            for member in item.split():
                for male_name in male_list:
                    if member.lower() == male_name.lower():
                        name_user_set.add(item)
                for female_name in female_list:
                    if member.lower() == female_name.lower():
                        name_user_set.add(item)
        if not name_user_set and local_list_name_block:
            name_user_set.add(local_list_name_block[0])
        return name_user_set
    except Exception as e:
        logging.error(f"Name extraction failed: {e}")
        return set()

def get_skills(file_content):
    try:
        skills_block = []
        found = False
        for line in file_content.split('\n'):
            if not found and match_section(line) == 'skills':
                found = True
                continue
            if found:
                if match_section(line) in SECTION_HEADERS:
                    break
                if line.strip():
                    skills_block.append(line.strip())
        text = ' '.join(skills_block)
        text = re.sub(r'[\u2022\u2022\-\n]', ',', text)
        candidates = [t.strip() for t in re.split(r'[;,]', text) if t.strip()]
        skill_set = set()
        for cand in candidates:
            key = cand.lower()
            for main_skill, aliases in skill_json_data.items():
                if key == main_skill.lower() or key in [a.lower() for a in aliases]:
                    skill_set.add(main_skill)
                    break
            else:
                if len(cand.split()) <= 3:
                    skill_set.add(cand.title())
        return skill_set
    except Exception as e:
        logging.error(f"Skills extraction failed: {e}")
        return set()

def get_education(file_content):
    try:
        section_lines = []
        collecting = False
        for line in file_content.split('\n'):
            if match_section(line) == 'education':
                collecting = True
                continue
            if collecting and match_section(line):
                break
            if collecting:
                section_lines.append(line.strip())
        dates = re.findall(regex_expression_date, '\n'.join(section_lines))
        institutions = [line for line in section_lines if re.search(regex_expression_school_college, line)]
        degrees = []
        for item in section_lines:
            for course in course_corpus_list:
                if course and course in item:
                    degrees.append(course)
        return {
            "education-dates": [d.strip() for d in dates],
            "school-college": [i.strip() for i in institutions],
            "education-degrees": [d.strip() for d in degrees]
        }
    except Exception as e:
        logging.error(f"Education extraction failed: {e}")
        return {"education-dates": [], "school-college": [], "education-degrees": []}

def get_work_experience(file_content):
    try:
        experience_block = []
        collecting = False
        for line in file_content.split("\n"):
            if match_section(line) == "experience":
                collecting = True
                continue
            if collecting and match_section(line):
                break
            if collecting:
                experience_block.append(line.strip())

        jobs = []
        for i, line in enumerate(experience_block):
            if re.search(regex_expression_date, line):
                role = ""
                company = ""
                details = ""

                # Backtrack for role and company
                for j in range(i - 1, max(i - 4, -1), -1):
                    if experience_block[j] and not re.search(regex_expression_date, experience_block[j]):
                        if not role:
                            role = experience_block[j]
                        elif not company:
                            company = experience_block[j]
                            break

                # Forward for details
                details = "\n".join(
                    line for line in experience_block[i + 1:i + 6]
                    if line and not re.search(regex_expression_date, line)
                )

                jobs.append({
                    "role": role.strip(),
                    "company": company.strip(),
                    "duration": line.strip(),
                    "details": details.strip()
                })

        return jobs
    except Exception as e:
        logging.error(f"Experience extraction failed: {e}")
        return []
# === Output Functions ===

def get_projects(file_content):
    try:
        projects = []
        collecting = False
        section_lines = []
        for line in file_content.split('\n'):
            if match_section(line) == 'projects':
                collecting = True
                continue
            if collecting and match_section(line):
                break
            if collecting:
                section_lines.append(line.strip())
        for line in section_lines:
            if not line:
                continue
            if any(keyword in line.lower() for keyword in ["python", "tool", "app", "system", "project"]):
                projects.append(line.strip())
        return projects
    except Exception as e:
        logging.error(f"Project extraction failed: {e}")
        return []

def pdf_to_text(filename):
    try:
        rsrcmgr = PDFResourceManager()
        retstr = StringIO()
        codec = 'utf-8'
        laparams = LAParams()
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
        with open(filename, 'rb') as fp:
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.get_pages(fp, check_extractable=True):
                interpreter.process_page(page)
        text = retstr.getvalue()
        device.close()
        retstr.close()
        return text
    except Exception as e:
        logging.error(f"pdf_to_text failed: {e}")
        return ""

def pdf_to_text_for_direct_upload(fp):
    try:
        rsrcmgr = PDFResourceManager()
        retstr = StringIO()
        codec = 'utf-8'
        laparams = LAParams()
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.get_pages(fp, check_extractable=True):
            interpreter.process_page(page)
        text = retstr.getvalue()
        device.close()
        retstr.close()
        return text
    except Exception as e:
        logging.error(f"pdf_to_text_for_direct_upload failed: {e}")
        return ""

def output_in_json(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list):
    try:
        data = {
            'emails': email_list or [],
            'name': list(name_list) if name_list else [],
            'phones': phone_list or [],
            'skills': list(skill_list) if skill_list else [],
            'experience': work_experience_dict or [],
            'education': education_dict or {},
            'projects': projects_list or []
        }
        return json.dumps(data, indent=4)
    except Exception as e:
        logging.error(f"Output JSON failed: {e}")
        return json.dumps({})

def output_in_xml(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list):
    if dicttoxml is None:
        raise ImportError("dicttoxml is not installed. Please install it to use XML output.")
    try:
        data = {
            'emails': email_list or [],
            'name': list(name_list) if name_list else [],
            'phones': phone_list or [],
            'skills': list(skill_list) if skill_list else [],
            'experience': work_experience_dict or [],
            'education': education_dict or {},
            'projects': projects_list or []
        }
        xml_bytes = dicttoxml.dicttoxml(data, custom_root='resume', attr_type=False)
        return parseString(xml_bytes).toprettyxml()
    except Exception as e:
        logging.error(f"Output XML failed: {e}")
        return "<resume></resume>"

def extract_file_content(filename, output_type='json'):
    try:
        file_content = ''
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        if ext == '.pdf':
            file_content = pdf_to_text(filename)
        elif ext == '.txt':
            file_content = open(filename, 'r', encoding='utf-8').read()
        elif ext == '.docx':
            doc = Document(filename)
            file_content = '\n'.join([p.text for p in doc.paragraphs])
        elif ext == '.rtf':
            file_content = rtf_to_text(open(filename, 'r').read())
        elif ext == '.html':
            soup = BeautifulSoup(open(filename, 'r', encoding='utf-8').read(), 'html.parser')
            file_content = soup.get_text()
        else:
            raise ValueError("Unsupported file type: " + ext)

        name_list = get_name(file_content)
        email_list = get_email(file_content)
        phone_list = get_phone_number(file_content)
        skill_list = get_skills(file_content)
        education_dict = get_education(file_content)
        work_experience_dict = get_work_experience(file_content)
        projects_list = get_projects(file_content)

        if output_type == 'json':
            return output_in_json(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list)
        else:
            return output_in_xml(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list)
    except Exception as e:
        logging.error(f"extract_file_content failed: {e}")
        return json.dumps({}) if output_type == 'json' else "<resume></resume>"

def read_file_content_directly(file_object, output_type='json'):
    try:
        file_content = ''
        filename = str(file_object)
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        if ext == '.txt':
            file_content = file_object.read().decode('utf-8')
        elif ext == '.docx':
            document = Document(file_object)
            file_content = '\n'.join(p.text for p in document.paragraphs)
        elif ext == '.pdf':
            fp = file_object.open('rb')
            file_content = pdf_to_text_for_direct_upload(fp)
        elif ext == '.rtf':
            file_content = rtf_to_text(file_object.read().decode('utf-8'))
        elif ext == '.html':
            soup = BeautifulSoup(file_object.read(), 'html.parser')
            file_content = soup.get_text()
        else:
            raise ValueError("Unsupported file type: " + ext)

        name_list = get_name(file_content)
        email_list = get_email(file_content)
        phone_list = get_phone_number(file_content)
        skill_list = get_skills(file_content)
        education_dict = get_education(file_content)
        work_experience_dict = get_work_experience(file_content)
        projects_list = get_projects(file_content)

        if output_type == 'json':
            return output_in_json(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list)
        else:
            return output_in_xml(email_list, name_list, phone_list, skill_list, education_dict, work_experience_dict, projects_list)
    except Exception as e:
        logging.error(f"read_file_content_directly failed: {e}")
        return json.dumps({}) if output_type == 'json' else "<resume></resume>"